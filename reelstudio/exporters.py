"""Exporters: turn a :class:`ReelProject` into downloadable artifacts.

Text/data formats (TXT, SRT, VTT, CSV, JSON) are pure-Python and always
available. PDF (reportlab) and DOCX (python-docx) are optional — if the library
is missing the helper returns ``None`` so the UI can hide that button instead of
crashing.
"""

from __future__ import annotations

import csv
import io
import json

from .models import ReelProject


# --- Time helpers ------------------------------------------------------------

def _ts(seconds: float, sep: str = ",") -> str:
    """Format seconds as HH:MM:SS,mmm (SRT) or HH:MM:SS.mmm (VTT)."""
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1000)
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


# --- Subtitles ---------------------------------------------------------------

def to_srt(project: ReelProject) -> str:
    blocks = []
    for i, sc in enumerate(project.scenes, start=1):
        text = sc.subtitle.replace("\n", "\n")
        blocks.append(
            f"{i}\n{_ts(sc.start)} --> {_ts(sc.end)}\n{text}\n"
        )
    return "\n".join(blocks).strip() + "\n"


def to_vtt(project: ReelProject) -> str:
    blocks = ["WEBVTT\n"]
    for i, sc in enumerate(project.scenes, start=1):
        blocks.append(
            f"{i}\n{_ts(sc.start, '.')} --> {_ts(sc.end, '.')}\n{sc.subtitle}\n"
        )
    return "\n".join(blocks).strip() + "\n"


def to_plain_subtitles(project: ReelProject) -> str:
    lines = []
    for sc in project.scenes:
        lines.append(f"[{_ts(sc.start, '.')}] {sc.subtitle.replace(chr(10), ' ')}")
    return "\n".join(lines) + "\n"


# --- Script / data -----------------------------------------------------------

def to_script_txt(project: ReelProject) -> str:
    out = io.StringIO()
    out.write(f"{project.title}\n{'=' * len(project.title)}\n\n")
    out.write(f"Topic: {project.topic}\n")
    out.write(f"Type: {project.content_type} | Duration: {project.duration}s | Tone: {project.tone}\n")
    out.write(f"Voice: {project.audio.voice_style} | CTA: {project.cta}\n\n")
    out.write(f"HOOK: {project.primary_hook}\n\n")
    out.write("FULL VOICEOVER SCRIPT\n---------------------\n")
    for sc in project.scenes:
        out.write(f"\n[Scene {sc.number} | {sc.role.upper()} | {sc.start:g}-{sc.end:g}s]\n")
        out.write(f"VO: {sc.voiceover}\n")
        out.write(f"On-screen: {sc.text_overlay}\n")
    return out.getvalue()


def to_csv_storyboard(project: ReelProject) -> str:
    out = io.StringIO()
    w = csv.writer(out)
    w.writerow([
        "Scene", "Role", "Start", "End", "Duration", "Visual Type", "Shot Size",
        "Camera", "Visual Description", "Background", "Text Overlay", "Subtitle",
        "Voiceover", "Sound Effect", "Transition", "Music Mood",
    ])
    for sc in project.scenes:
        w.writerow([
            sc.number, sc.role, sc.start, sc.end, sc.duration, sc.visual_type,
            sc.shot_size, sc.camera_movement, sc.visual_description, sc.background,
            sc.text_overlay, sc.subtitle.replace("\n", " "), sc.voiceover,
            sc.sound_effect, sc.transition, sc.music_mood,
        ])
    return out.getvalue()


def to_json(project: ReelProject) -> str:
    return json.dumps(project.to_dict(), indent=2, ensure_ascii=False)


# --- Editor briefs (markdown text) ------------------------------------------

def _header(project: ReelProject) -> str:
    return (
        f"# {project.title}\n\n"
        f"- Topic: {project.topic}\n"
        f"- Type: {project.content_type} | Duration: {project.duration}s "
        f"| Total scenes: {len(project.scenes)}\n"
        f"- Tone: {project.tone} | Voice: {project.audio.voice_style}\n"
        f"- Color palette: {project.color_palette}\n\n"
    )


def canva_brief(project: ReelProject) -> str:
    out = [_header(project)]
    out.append("## Canva Reel Brief\n")
    out.append("- Design size: **Instagram Reel 1080 × 1920**\n")
    out.append(f"- Number of pages/scenes: **{len(project.scenes)}**\n")
    out.append(f"- Color palette: {project.color_palette}\n")
    out.append("- Suggested fonts: heavy display for hooks (Anton / Montserrat ExtraBold), clean sans for body.\n\n")
    for sc in project.scenes:
        out.append(
            f"### Page {sc.number} — {sc.role.upper()} ({sc.duration:g}s)\n"
            f"- Background: {sc.background}\n"
            f"- Text on page: **{sc.text_overlay}** / caption: \"{sc.subtitle.splitlines()[0]}\"\n"
            f"- Animation: {project.motion_graphics['entry_animation']} in, "
            f"{project.motion_graphics['exit_animation']} out\n"
            f"- Element/icon: {', '.join(sc.stock.icon[:3])}\n"
            f"- Upload assets needed: {sc.visual_type.lower()} for this page\n\n"
        )
    out.append("**Final export:** Share → Download → MP4 Video, 1080 × 1920.\n")
    return "".join(out)


def capcut_brief(project: ReelProject) -> str:
    out = [_header(project)]
    out.append("## CapCut Editing Brief\n\n### Timeline\n")
    for sc in project.scenes:
        out.append(
            f"- Clip {sc.number} [{sc.start:g}s–{sc.end:g}s, {sc.duration:g}s] — {sc.visual_type}\n"
            f"  - Text overlay: **{sc.text_overlay}**\n"
            f"  - Subtitle ({_ts(sc.start, '.')}→{_ts(sc.end, '.')}): {sc.subtitle.replace(chr(10),' ')}\n"
            f"  - Transition out: {sc.transition}\n"
            f"  - SFX: {sc.sound_effect}\n"
        )
    out.append(
        "\n### Effects / filters\n"
        "- Auto Captions (then restyle to bold, bottom-safe).\n"
        "- Subtle 'Zoom' or 'Shake' on the hook; clean look filter elsewhere.\n"
        "- Beat markers: tap to beat on the music track, cut on the beat.\n\n"
        "### Music & beat points\n"
        f"- Mood: {project.audio.music_mood}\n"
        f"- Search: {', '.join(project.audio.music_keywords)}\n"
        "- Place cuts on detected beats; drop SFX on emphasis words.\n\n"
        "### Export settings\n"
        "- Format: MP4 | Aspect: 9:16 | Resolution: 1080 × 1920 | FPS: 30\n"
        f"- Recommended length: ~{project.duration}s\n"
    )
    return "".join(out)


def premiere_brief(project: ReelProject) -> str:
    out = [_header(project)]
    out.append(
        "## Premiere Pro Brief\n\n"
        "- Sequence: 1080×1920, 30fps, square pixels.\n"
        "- Tracks: V1 footage/stills, V2 captions (Essential Graphics), "
        "A1 voiceover, A2 music, A3 SFX.\n\n### EDL\n"
    )
    for sc in project.scenes:
        out.append(
            f"- {sc.number:02d}. {_ts(sc.start, '.')}–{_ts(sc.end, '.')} | {sc.visual_type} | "
            f"{sc.shot_size} | {sc.camera_movement} | caption: \"{sc.text_overlay}\" | "
            f"transition: {sc.transition}\n"
        )
    out.append(
        "\n- Captions: use Essential Graphics → Captions, import the SRT for auto-timed text.\n"
        "- Export: H.264, 1080×1920, 30fps, ~10-12 Mbps, AAC audio.\n"
    )
    return "".join(out)


def finalcut_brief(project: ReelProject) -> str:
    out = [_header(project)]
    out.append(
        "## Final Cut Pro Brief\n\n"
        "- Project: Vertical 1080×1920, 30fps.\n"
        "- Roles: Video, Titles, Dialogue (VO), Music, Effects (SFX).\n\n### Storyline\n"
    )
    for sc in project.scenes:
        out.append(
            f"- Clip {sc.number}: {sc.duration:g}s | {sc.visual_type} | {sc.shot_size} | "
            f"title \"{sc.text_overlay}\" | {sc.transition} | SFX: {sc.sound_effect}\n"
        )
    out.append(
        "\n- Import SRT as captions; use Custom Overlay titles for hooks.\n"
        "- Share → Master File (or 'Prepare for Instagram'), 1080×1920, 30fps, H.264/HEVC.\n"
    )
    return "".join(out)


def shot_list(project: ReelProject) -> str:
    out = [_header(project), "## Shot List\n\n"]
    for sc in project.scenes:
        out.append(
            f"- Shot {sc.number} | {sc.shot_size} | {sc.camera_movement} | {sc.visual_type}\n"
            f"  {sc.visual_description}\n"
        )
    return "".join(out)


def asset_checklist(project: ReelProject) -> str:
    talking = [s.number for s in project.scenes if s.visual_type == "Talking head"]
    broll = [s.number for s in project.scenes if s.visual_type in ("B-roll", "Stock video")]
    ai_img = [s.number for s in project.scenes if s.needs_ai_image]
    ai_vid = [s.number for s in project.scenes if s.needs_ai_video]
    screen = [s.number for s in project.scenes if s.visual_type == "Screen recording"]
    anim = [s.number for s in project.scenes if s.visual_type in ("Animation", "Text-only motion graphic")]

    def line(label: str, scenes: list[int]) -> str:
        mark = "x" if scenes else " "
        detail = f"scenes {', '.join(map(str, scenes))}" if scenes else "none"
        return f"- [{mark}] {label} ({detail})\n"

    out = [_header(project), "## Multimedia Asset Checklist\n\n"]
    out.append(line("Talking-head video clips", talking))
    out.append(line("B-roll / stock clips", broll))
    out.append(line("AI images", ai_img))
    out.append(line("AI video clips", ai_vid))
    out.append(line("Screen recordings", screen))
    out.append(line("Animations / motion graphics", anim))
    out.append("- [ ] Icons (per scene — see storyboard)\n")
    out.append("- [ ] Backgrounds\n")
    out.append(f"- [ ] Music ({project.audio.music_mood}, royalty-free)\n")
    out.append("- [ ] Sound effects (per scene)\n")
    out.append("- [ ] Logo / brand assets\n")
    out.append(f"- [ ] Thumbnail image (\"{project.thumbnail.text}\")\n")
    out.append("- [ ] Captions / subtitles (SRT + VTT)\n")
    return "".join(out)


def full_brief_markdown(project: ReelProject) -> str:
    """One combined human-readable production brief (used for PDF/DOCX too)."""
    parts = [
        _header(project),
        "## Strategy\n",
        f"- Goal: {project.strategy.goal}\n- Audience: {project.strategy.audience}\n"
        f"- Angle: {project.strategy.angle}\n- Platform notes: {project.strategy.platform_notes}\n"
        f"- Hashtags: {' '.join(project.strategy.hashtags)}\n\n",
        "## Hook options\n" + "".join(f"- {h}\n" for h in project.hook_options) + "\n",
        "## Storyboard\n",
    ]
    for sc in project.scenes:
        parts.append(
            f"### Scene {sc.number} — {sc.role.upper()} ({sc.start:g}–{sc.end:g}s)\n"
            f"- Visual: {sc.visual_type} | {sc.shot_size} | {sc.camera_movement}\n"
            f"- Description: {sc.visual_description}\n"
            f"- Overlay: {sc.text_overlay}\n- VO: {sc.voiceover}\n- Subtitle: {sc.subtitle.replace(chr(10),' ')}\n"
            f"- SFX: {sc.sound_effect} | Transition: {sc.transition} | Music: {sc.music_mood}\n"
        )
        if sc.image_prompt:
            parts.append(f"- AI image prompt: {sc.image_prompt.as_text()}\n")
        if sc.video_prompt:
            parts.append(f"- AI video prompt: {sc.video_prompt.as_text()}\n")
        parts.append("\n")
    parts.append(
        "## Audio plan\n"
        f"- Voice: {project.audio.voice_style} ({project.audio.emotional_tone}), {project.audio.voice_speed}\n"
        f"- Music: {project.audio.music_mood} — search: {', '.join(project.audio.music_keywords)}\n"
        f"- Silence: {' '.join(project.audio.silence_moments)}\n\n"
    )
    parts.append(
        "## Thumbnail\n"
        f"- Text: {project.thumbnail.text}\n- Layout: {project.thumbnail.layout}\n"
        f"- AI prompt: {project.thumbnail.ai_prompt}\n\n"
    )
    parts.append("## Safety & quality notes\n" + "".join(f"- {n}\n" for n in project.safety_notes))
    return "".join(parts)


# --- Optional binary formats -------------------------------------------------

def to_pdf(project: ReelProject) -> bytes | None:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except Exception:
        return None
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=project.title)
    styles = getSampleStyleSheet()
    flow = []
    for raw in full_brief_markdown(project).splitlines():
        line = raw.rstrip()
        if not line:
            flow.append(Spacer(1, 6))
            continue
        if line.startswith("### "):
            flow.append(Paragraph(line[4:], styles["Heading3"]))
        elif line.startswith("## "):
            flow.append(Paragraph(line[3:], styles["Heading2"]))
        elif line.startswith("# "):
            flow.append(Paragraph(line[2:], styles["Title"]))
        else:
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flow.append(Paragraph(safe.lstrip("- "), styles["Normal"]))
    doc.build(flow)
    return buf.getvalue()


def to_docx(project: ReelProject) -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None
    doc = Document()
    for raw in full_brief_markdown(project).splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            doc.add_paragraph(line.lstrip("- "))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
